import os
import re
from datetime import datetime
from docx import Document
from duckduckgo_search import DDGS
from bs4 import BeautifulSoup
import requests
try:
    from .llm import query_llm
    from . import prompts
except ImportError:
    from llm import query_llm
    import prompts
import PyPDF2
import io
import json


def save_output(report, code, execution_result, timestamp, iteration):
    # timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_dir = f"./output_agent/output_{timestamp}"
    os.makedirs(output_dir, exist_ok=True)

    # save report
    if report:
        doc = Document()
        doc.add_heading(f"Research Report - Iteration {iteration + 1}", level=1)
        doc.add_paragraph(report)
        report_file = os.path.join(output_dir, f"research_report_iteration_{iteration + 1}.docx")
        doc.save(report_file)

    # save code
    if code:
        code_file = os.path.join(output_dir, f"code_iteration_{iteration + 1}.py")
        with open(code_file, "w") as f:
            f.write(code)

    # save execution results
    if execution_result: 
        result_file = os.path.join(output_dir, f"execution_result_{iteration + 1}.txt")
        with open(result_file, "w") as f:
            f.write(execution_result)

    print(f"\nOutputs saved for iteration {iteration + 1} in {output_dir}")

# function to clean up the report to conform to professional standards
def clean_report(text):
    """
    Removes LLM reasoning, markdown, and formatting artifacts from report output.
    """
    # Remove <think>...</think> blocks
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)

    # Remove markdown-style headers and separators
    text = re.sub(r"^\s*#+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*-{3,}\s*$", "", text, flags=re.MULTILINE)

    # Remove extra leading/trailing whitespace
    return text.strip()
    
def extract_code_only(text):
    """
    Extracts clean Python code from an LLM response by:
    - Removing <think>...</think> blocks
    - Removing markdown code fences (```python ... ```)
    - Returning only executable code with inline comments
    """
    # Remove <think>...</think> sections
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL)

    # Extract the code inside ```python ... ``` blocks, if they exist
    match = re.search(r"```(?:python)?\n(.*?)```", text, re.DOTALL)
    code = match.group(1).strip() if match else text.strip()

    return code
    
def quick_duckduckgo_search(query, max_results=3):
    print(f"Performing quick DuckDuckGo search for: '{query}'")
    try:
        with DDGS() as ddgs:
            results = ddgs.text(query)
            top_results = list(results)[:max_results]

        raw_text = "\n\n".join(
            f"{i+1}. {r['title']}\n    {r['href']}\n    {r.get('body', '').strip()}"
            for i, r in enumerate(top_results)
        )

        print("Summarizing top search results...\n")

        summary_prompt = prompts.get_quick_search_summary_prompt(query, raw_text)

        raw_summary = query_llm(summary_prompt).strip()
        # Remove <think>...</think> block if present
        summary = re.sub(r"<think>.*?</think>", "", raw_summary, flags=re.DOTALL).strip()

        return f"Answer:\n{summary}\n\nBased on DuckDuckGo Search Results:\n\n{raw_text}"

    except Exception as e:
        return f"DuckDuckGo search failed: {e}"


def extract_pdf_text(pdf_path):
    """
    Extract text from a PDF file using multiple methods for better coverage.
    Returns a dictionary with filename and extracted text.
    """
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        
        return {
            "filename": os.path.basename(pdf_path),
            "content": text.strip()
        }
    except Exception as e:
        print(f"PyPDF2 failed for {pdf_path}: {e}")
        return {
            "filename": os.path.basename(pdf_path),
            "content": f"Error extracting text from PDF: {e}"
        }


def process_pdfs(pdf_paths):
    """
    Process multiple PDF files and return their extracted text.
    """
    if not pdf_paths:
        return ""
    
    pdf_contents = []
    for pdf_path in pdf_paths:
        if os.path.exists(pdf_path):
            content = extract_pdf_text(pdf_path)
            pdf_contents.append(content)
        else:
            print(f"Warning: PDF file not found: {pdf_path}")
    
    if not pdf_contents:
        return ""
    
    # Format PDF contents for inclusion in prompts
    formatted_pdfs = "\n\n".join(
        f"PDF: {content['filename']}\n"
        f"{'='*50}\n"
        f"{content['content']}\n"
        f"{'='*50}"
        for content in pdf_contents
    )
    
    return formatted_pdfs


# def process_pdfs(pdf_paths):
#     """
#     Process multiple PDF files and return their extracted text.
#     """
#     if not pdf_paths:
#         return ""
    
#     pdf_contents = []
#     # for pdf_path in pdf_paths:
#     #     if os.path.exists(pdf_path):
#     #         content = extract_pdf_text(pdf_path)
#     #         pdf_contents.append(content)
#     #     else:
#     #         print(f"Warning: PDF file not found: {pdf_path}")

#     def iter_pdf_files(paths):
#         for path in paths:
#             if not os.path.exists(path):
#                 print(f"Warning: PDF file not found: {path}")
#                 continue
#             if os.path.isdir(path):
#                 for root, _, files in os.walk(path):
#                     for filename in sorted(files):
#                         if filename.lower().endswith('.pdf'):
#                             yield os.path.join(root, filename)
#             else:
#                 yield path

#     for pdf_path in iter_pdf_files(pdf_paths):
#         content = extract_pdf_text(pdf_path)
#         pdf_contents.append(content)


    
#     if not pdf_contents:
#         return ""
    
#     # Format PDF contents for inclusion in prompts
#     formatted_pdfs = "\n\n".join(
#         f"PDF: {content['filename']}\n"
#         f"{'='*50}\n"
#         f"{content['content']}\n"
#         f"{'='*50}"
#         for content in pdf_contents
#     )
    
#     return formatted_pdfs


def process_links(link_paths):
    """
    Process multiple links and return their extracted content.
    """
    if not link_paths:
        return ""
    
    link_contents = []
    for link in link_paths:
        content = extract_link_content(link)
        if content:
            link_contents.append({
                "url": link,
                "content": content
            })
        else:
            print(f"Warning: Could not extract content from link: {link}")
    
    if not link_contents:
        return ""
    
    # Format link contents for inclusion in prompts
    formatted_links = "\n\n".join(
        f"Link: {content['url']}\n"
        f"{'='*50}\n"
        f"{content['content']}\n"
        f"{'='*50}"
        for content in link_contents
    )
    
    return formatted_links


def extract_link_content(url):
    """
    Extract content from a URL using various methods.
    """
    try:
        # Handle HuggingFace URLs
        if 'huggingface.co' in url and '/blob/' in url:
            return extract_huggingface_content(url)
        
        # Handle GitHub URLs
        elif 'github.com' in url and '/blob/' in url:
            return extract_github_content(url)
        
        # Handle other URLs (basic web scraping)
        else:
            return extract_basic_content(url)
            
    except Exception as e:
        print(f"Error extracting content from {url}: {e}")
        return None


def extract_huggingface_content(url):
    """Extract content from HuggingFace URLs"""
    try:
        # Convert blob URL to resolve URL
        resolve_url = url.replace('/blob/', '/resolve/')
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(resolve_url, headers=headers, timeout=10)
        if response.status_code == 200:
            content = response.text
            
            # Try to parse as Jupyter notebook
            if url.endswith('.ipynb'):
                return parse_jupyter_notebook(content)
            else:
                return content[:2000] + "..." if len(content) > 2000 else content
        else:
            print(f"Failed to fetch HuggingFace content: {response.status_code}")
            return None
            
    except Exception as e:
        print(f"Error fetching HuggingFace content: {e}")
        return None


def extract_github_content(url):
    """Extract content from GitHub URLs"""
    try:
        # Convert blob URL to raw URL
        raw_url = url.replace('github.com', 'raw.githubusercontent.com').replace('/blob/', '/')
        
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(raw_url, headers=headers, timeout=10)
        if response.status_code == 200:
            content = response.text
            return content[:2000] + "..." if len(content) > 2000 else content
        else:
            print(f"Failed to fetch GitHub content: {response.status_code}")
            return None
            
    except Exception as e:
        print(f"Error fetching GitHub content: {e}")
        return None


def extract_basic_content(url):
    """Extract basic web content"""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        
        response = requests.get(url, headers=headers, timeout=10)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, 'html.parser')
            text_content = soup.get_text()
            return text_content[:2000] + "..." if len(text_content) > 2000 else text_content
        else:
            print(f"Failed to fetch basic content: {response.status_code}")
            return None
            
    except Exception as e:
        print(f"Error fetching basic content: {e}")
        return None


def parse_jupyter_notebook(content):
    """Parse Jupyter notebook content"""
    try:
        notebook = json.loads(content)
        cells = notebook.get('cells', [])
        
        extracted_text = []
        for i, cell in enumerate(cells):
            cell_type = cell.get('cell_type', '')
            source = cell.get('source', [])
            
            if isinstance(source, list):
                text = ''.join(source)
            else:
                text = str(source)
            
            if cell_type == 'markdown':
                extracted_text.append(f"## Cell {i+1} (Markdown)\n{text}\n")
            elif cell_type == 'code':
                extracted_text.append(f"## Cell {i+1} (Code)\n```python\n{text}\n```\n")
        
        return '\n'.join(extracted_text)
    except Exception as e:
        print(f"Error parsing notebook: {e}")
        return content


def explore_files_directory(directory_path):
    """
    Explore a directory and return a simple file listing for LLM analysis.
    """
    try:
        if not os.path.exists(directory_path):
            return f"Error: Directory {directory_path} does not exist."
        
        file_list = []
        
        # Walk through the directory and collect file information
        for root, dirs, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                relative_path = os.path.relpath(file_path, directory_path)
                
                # Get file size
                try:
                    file_size = os.path.getsize(file_path)
                    size_str = f"{file_size:,} bytes"
                    if file_size > 1024*1024:
                        size_str = f"{file_size/(1024*1024):.1f} MB"
                    elif file_size > 1024:
                        size_str = f"{file_size/1024:.1f} KB"
                except:
                    size_str = "Unknown size"
                
                file_list.append({
                    'path': relative_path,
                    'size': size_str,
                    'extension': os.path.splitext(file)[1].lower()
                })
        
        # Create a simple report for LLM analysis
        report = f"FILES DIRECTORY EXPLORATION\n"
        report += f"Directory: {directory_path}\n"
        report += f"Total files found: {len(file_list)}\n\n"
        report += f"FILE LISTING:\n"
        report += "-" * 50 + "\n"
        
        for file_info in file_list:
            report += f"{file_info['path']} ({file_info['size']})\n"
        
        return report
        
    except Exception as e:
        return f"Error exploring files directory: {str(e)}"